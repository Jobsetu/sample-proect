from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from bytez import Bytez
import os
from pdfminer.high_level import extract_text
from docx import Document as DocxDocument

app = Flask(__name__)
CORS(app)

# Initialize Bytez SDK
# Ideally, get key from environment variable
BYTEZ_KEY = os.environ.get("BYTEZ_API_KEY", "e7bcd604f04b496ca11602337f3a81fc")
sdk = Bytez(BYTEZ_KEY)
model = sdk.model("google/gemini-2.5-pro")

@app.route('/api/upload-resume', methods=['POST'])
def upload_resume():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
            
        text = ""
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            # Save temporarily to parse
            temp_path = f"temp_{file.filename}"
            file.save(temp_path)
            try:
                text = extract_text(temp_path)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
        elif filename.endswith('.docx'):
            doc = DocxDocument(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            
        else:
            return jsonify({'error': 'Unsupported file format. Please upload PDF or DOCX'}), 400
            
        if not text.strip():
            return jsonify({'error': 'Could not extract text from file'}), 400
            
        # Use AI to parse the text into our JSON structure
        print(f"[UPLOAD] Extracted {len(text)} chars. Parsing with AI...")
        
        prompt = f"""
        You are a resume parser. Extract the following information from the resume text below and return it as a JSON object.
        
        Resume Text:
        {text[:10000]}  # Limit text length to avoid token limits
        
        Required JSON Structure:
        {{
            "personalInfo": {{
                "name": "Full Name",
                "email": "email@example.com",
                "phone": "Phone Number",
                "location": "City, State",
                "linkedin": "LinkedIn URL",
                "github": "GitHub URL",
                "title": "Current Job Title"
            }},
            "sections": [
                {{
                    "id": "summary",
                    "title": "Professional Summary",
                    "content": "Summary text..."
                }},
                {{
                    "id": "experience",
                    "title": "Experience",
                    "items": [
                        {{
                            "company": "Company Name",
                            "role": "Job Title",
                            "location": "Location",
                            "startDate": "Start Date",
                            "endDate": "End Date",
                            "bullets": ["Achievement 1", "Achievement 2"]
                        }}
                    ]
                }},
                {{
                    "id": "education",
                    "title": "Education",
                    "items": [
                        {{
                            "school": "University Name",
                            "degree": "Degree Name",
                            "field": "Field of Study",
                            "graduationDate": "Year"
                        }}
                    ]
                }},
                {{
                    "id": "skills",
                    "title": "Skills",
                    "items": ["Skill 1", "Skill 2"]
                }},
                {{
                    "id": "projects",
                    "title": "Projects",
                    "items": [
                        {{
                            "title": "Project Name",
                            "description": "Project Description",
                            "technologies": ["Tech 1", "Tech 2"]
                        }}
                    ]
                }}
            ]
        }}
        
        Return ONLY valid JSON. Do not include markdown formatting.
        """
        
        response = model.run([{"role": "user", "content": prompt}])
        
        # Extract and clean output
        if hasattr(response, 'output') and isinstance(response.output, dict):
            output = response.output.get('content', '')
        elif isinstance(response, dict) and 'output' in response:
            output = response['output'].get('content', '')
        elif hasattr(response, 'content'):
            output = response.content
        else:
            output = str(response)
            
        # Clean markdown
        if output.startswith('```'):
            output = output.split('\n', 1)[1]
            if output.endswith('```'):
                output = output.rsplit('\n', 1)[0]
        
        # Parse and sanitize JSON
        import json
        try:
            parsed_json = json.loads(output)
            
            def sanitize_skills_list(items):
                clean_items = []
                if isinstance(items, list):
                    for item in items:
                        if isinstance(item, str):
                            clean_items.append(item)
                        elif isinstance(item, dict):
                            # Extract ALL string values from the dict recursively
                            for key, val in item.items():
                                if isinstance(val, str):
                                    clean_items.append(val)
                                elif isinstance(val, list):
                                    clean_items.extend(sanitize_skills_list(val))
                                elif isinstance(val, dict):
                                    # Handle nested dicts (e.g. { "skill": { "name": "Java" } })
                                    clean_items.extend(sanitize_skills_list([val]))
                        elif isinstance(item, list):
                            clean_items.extend(sanitize_skills_list(item))
                        else:
                            # Fallback for numbers/booleans -> stringify
                            if item is not None:
                                clean_items.append(str(item))
                return clean_items

            # Sanitize skills to ensure they are strings
            if 'sections' in parsed_json:
                for section in parsed_json['sections']:
                    if section.get('id') == 'skills' and 'items' in section:
                        original_items = section['items']
                        cleaned = sanitize_skills_list(original_items)
                        print(f"[SANITIZER] Skills cleaned. Original count: {len(original_items)}, New count: {len(cleaned)}")
                        if len(cleaned) > 0:
                             print(f"[SANITIZER] Sample: {cleaned[:3]}")
                        section['items'] = cleaned
            
            output = json.dumps(parsed_json)
        except Exception as e:
            print(f"[SANITIZER] Error: {e}")
            # Even if parsing fails, we might want to try to salvage something, but for now return raw
            pass 
        
        return jsonify({'output': output})

    except Exception as e:
        print(f"Upload error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    try:
        from resume_service import ResumeService
        
        data = request.json
        user_input = data.get('input', '')
        user_resume = data.get('resume', {})
        job_description = data.get('jobDescription', '')
        job_title = data.get('jobTitle', 'Position')
        company_name = data.get('companyName', 'Company')
        
        print(f"[RESUME] Generating resume for {job_title} at {company_name}")
        
        # Parse user resume
        if not user_resume or not isinstance(user_resume, dict):
            print("[RESUME] No valid resume data provided, creating minimal resume structure")
            user_resume = {
                'personalInfo': {
                    'name': 'Professional',
                    'email': '',
                    'phone': '',
                    'location': ''
                },
                'sections': []
            }
        parsed_resume = ResumeService.parse_user_resume(user_resume)
        
        # Try AI generation first
        try:
            # Construct a better prompt that explicitly asks for skills from both sources
            prompt_content = f"""
            You are an expert resume writer. Create a tailored resume for the following job.
            
            JOB DESCRIPTION:
            {job_description}
            
            CANDIDATE PROFILE:
            {json.dumps(parsed_resume)}
            
            USER INSTRUCTIONS:
            {user_input}
            
            REQUIREMENTS:
            1. Extract relevant skills from BOTH the Job Description and the Candidate Profile.
            2. The 'skills' section MUST be a simple list of strings (e.g., ["Python", "React"]). NO OBJECTS.
            3. Tailor the Professional Summary to match the job.
            4. Highlight relevant experience.
            
            Return ONLY valid JSON in the following structure:
            {{
                "personalInfo": {{ ... }},
                "sections": [
                    {{ "id": "summary", "title": "Professional Summary", "content": "..." }},
                    {{ "id": "skills", "title": "Skills", "items": ["Skill1", "Skill2"] }},
                    {{ "id": "experience", "title": "Experience", "items": [ ... ] }},
                    {{ "id": "education", "title": "Education", "items": [ ... ] }},
                    {{ "id": "projects", "title": "Projects", "items": [ ... ] }}
                ]
            }}
            """
            
            print(f"[RESUME] Calling AI model with input length: {len(prompt_content)}")
            response = model.run([
                {
                    "role": "user",
                    "content": prompt_content
                }
            ])
            
            # Extract text output
            if hasattr(response, 'output') and isinstance(response.output, dict):
                text_output = response.output.get('content', '')
            elif isinstance(response, dict) and 'output' in response:
                text_output = response['output'].get('content', '')
            elif hasattr(response, 'content'):
                text_output = response.content
            else:
                text_output = str(response)
            
            # Clean up the output - remove any markdown code blocks if present
            if text_output:
                text_output = text_output.strip()
                # Remove markdown code blocks if the AI wrapped it
                if text_output.startswith('```'):
                    lines = text_output.split('\n')
                    # Remove first line (```markdown or ```)
                    if lines[0].startswith('```'):
                        lines = lines[1:]
                    # Remove last line (```)
                    if lines and lines[-1].strip() == '```':
                        lines = lines[:-1]
                    text_output = '\n'.join(lines).strip()
            
            # Sanitize generated JSON
            try:
                gen_json = json.loads(text_output)
                
                def sanitize_skills_list(items):
                    clean_items = []
                    if isinstance(items, list):
                        for item in items:
                            if isinstance(item, str):
                                clean_items.append(item)
                            elif isinstance(item, dict):
                                for key, val in item.items():
                                    if isinstance(val, str):
                                        clean_items.append(val)
                                    elif isinstance(val, list):
                                        clean_items.extend(sanitize_skills_list(val))
                                    elif isinstance(val, dict):
                                        clean_items.extend(sanitize_skills_list([val]))
                            elif isinstance(item, list):
                                clean_items.extend(sanitize_skills_list(item))
                            else:
                                if item is not None:
                                    clean_items.append(str(item))
                    return clean_items

                if 'sections' in gen_json:
                    for section in gen_json['sections']:
                        if section.get('id') == 'skills' and 'items' in section:
                            original_items = section['items']
                            cleaned = sanitize_skills_list(original_items)
                            print(f"[SANITIZER-GEN] Skills cleaned. Original count: {len(original_items)}, New count: {len(cleaned)}")
                            section['items'] = cleaned
                
                text_output = json.dumps(gen_json)
            except Exception as e:
                print(f"[RESUME] JSON sanitization failed: {e}")

            # Validate that it looks like a resume (starts with # or contains resume sections)
            # Also check if it's a guide/template (contains words like "template", "example", "here is")
                
                # Validate that it looks like a resume (starts with # or contains resume sections)
                # Also check if it's a guide/template (contains words like "template", "example", "here is")
                text_lower = text_output.lower()
                is_guide = any(word in text_lower[:500] for word in ['template', 'example resume', 'here is a', 'tips for', 'how to', 'guide to', 'sample resume'])
                
                is_valid_resume = (
                    text_output and 
                    len(text_output) > 200 and 
                    not is_guide and
                    (text_output.strip().startswith('#') or 
                     'PROFESSIONAL SUMMARY' in text_output.upper() or
                     'EXPERIENCE' in text_output.upper() or
                     'EDUCATION' in text_output.upper())
                )
                
                if is_valid_resume:
                    print(f"[RESUME] AI generation successful, output length: {len(text_output)}")
                    return jsonify({"output": text_output, "source": "ai"})
                else:
                    print(f"[RESUME] AI output doesn't look like a resume (is_guide={is_guide}), using fallback. Output preview: {text_output[:300]}")
        
        except Exception as ai_error:
            print(f"[RESUME] AI generation failed: {ai_error}")
            import traceback
            traceback.print_exc()
        
        # Fallback to template-based generation
        print("[RESUME] Using fallback template generation")
        markdown_output = ResumeService.generate_detailed_resume_markdown(
            resume=parsed_resume,
            job_description=job_description or user_input,
            job_title=job_title,
            company_name=company_name
        )
        
        return jsonify({
            "output": markdown_output,
            "source": "template",
            "match_score": ResumeService.calculate_match_score(
                parsed_resume,
                ResumeService.extract_job_keywords(job_description or user_input)
            )
        })

    except Exception as e:
        print(f"[RESUME] Critical error: {e}")
        import traceback
        traceback.print_exc()
        
        # Last resort fallback
        return jsonify({
            "output": f"""# Professional Resume

**Email:** user@example.com | **Phone:** (555) 123-4567

## PROFESSIONAL SUMMARY

Experienced professional seeking opportunities in {data.get('jobTitle', 'technology')}.

## SKILLS

Python, JavaScript, React, Node.js, SQL, AWS, Docker, Git

## EXPERIENCE

**Software Engineer** | **Technology Company** | 2020 - Present

- Developed and maintained production systems
- Collaborated with cross-functional teams
- Implemented best practices and code reviews

## EDUCATION

**Bachelor of Science in Computer Science**
University of Technology | 2020
""",
            "source": "emergency_fallback"
        }), 200

@app.route('/api/generate-cover-letter', methods=['POST'])
def generate_cover_letter():
    try:
        data = request.json
        user_input = data.get('input')
        
        if not user_input:
            return jsonify({"error": "No input provided"}), 400

        response = model.run([
            {
                "role": "user",
                "content": f"Generate a professional cover letter based on the following details:\n{user_input}"
            }
        ])

        if hasattr(response, 'output') and isinstance(response.output, dict):
            text_output = response.output.get('content', '')
        elif isinstance(response, dict) and 'output' in response:
            text_output = response['output'].get('content', '')
        else:
            text_output = str(response)

        return jsonify({"output": text_output})

    except Exception as e:
        print(f"Exception: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/mock-interview', methods=['POST'])
def mock_interview():
    try:
        data = request.json
        messages = data.get('messages', [])
        
        if not messages:
             # Initial greeting
             messages = [{"role": "user", "content": "Start a mock interview for a software engineering role."}]

        response = model.run(messages)

        if hasattr(response, 'output') and isinstance(response.output, dict):
            text_output = response.output.get('content', '')
        elif isinstance(response, dict) and 'output' in response:
            text_output = response['output'].get('content', '')
        else:
            text_output = str(response)

        return jsonify({"output": text_output})

    except Exception as e:
        print(f"Exception: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        data = request.json
        resume = data.get('resume', {})
        
        doc = Document()
        
        # Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        # Header
        personal = resume.get('personalInfo', {})
        name = personal.get('name', 'Your Name')
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = header.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        
        contact_info = []
        if personal.get('email'): contact_info.append(personal['email'])
        if personal.get('phone'): contact_info.append(personal['phone'])
        if personal.get('location'): contact_info.append(personal['location'])
        if personal.get('linkedin'): contact_info.append(personal['linkedin'])
        
        if contact_info:
            contact_p = doc.add_paragraph(' | '.join(contact_info))
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_p.paragraph_format.space_after = Pt(12)

        # Sections
        for section in resume.get('sections', []):
            # Title
            doc.add_heading(section.get('title', '').upper(), level=2)
            
            # Content
            if section.get('id') == 'summary' or section.get('content'):
                doc.add_paragraph(section.get('content', ''))
            
            elif section.get('id') == 'skills':
                items = section.get('items', [])
                if isinstance(items, list):
                    doc.add_paragraph(', '.join(items))
                else:
                    doc.add_paragraph(str(items))
            
            elif section.get('items'):
                for item in section.get('items', []):
                    # Role/Title
                    p = doc.add_paragraph()
                    title = item.get('role') or item.get('title') or item.get('degree') or ''
                    company = item.get('company') or item.get('school') or item.get('subtitle') or ''
                    date = item.get('startDate') or item.get('year') or ''
                    if item.get('endDate'): date += f" - {item.get('endDate')}"
                    
                    run = p.add_run(title)
                    run.bold = True
                    if company: p.add_run(f" | {company}")
                    if date: p.add_run(f" | {date}")
                    
                    # Bullets
                    if item.get('bullets'):
                        for bullet in item.get('bullets', []):
                            doc.add_paragraph(bullet, style='List Bullet')
                    
                    # Description
                    if item.get('description'):
                        doc.add_paragraph(item.get('description'))

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0, 2)
        size = buffer.tell()
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name='resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            max_age=0
        )

    except Exception as e:
        print(f"DOCX Generation Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/download-cover-letter-docx', methods=['POST'])
def download_cover_letter_docx():
    try:
        from docx import Document
        from docx.shared import Pt
        import io

        data = request.json
        cover_letter_text = data.get('text', '')
        
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Sanitize text to remove characters that might break XML
        def sanitize(text):
            if not text: return ""
            # Keep only printable characters and newlines
            return "".join(ch for ch in text if ch == '\n' or (ord(ch) >= 32 and ord(ch) != 127))

        # Split by newlines and add paragraphs
        for line in cover_letter_text.split('\n'):
            clean_line = sanitize(line.strip())
            if clean_line:
                doc.add_paragraph(clean_line)
            else:
                # Add empty paragraph for spacing
                doc.add_paragraph('')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0, 2) # Seek to end
        size = buffer.tell()
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name='cover_letter.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            max_age=0
        )

    except Exception as e:
        print(f"Cover Letter DOCX Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "service": "job-yatra-backend"}), 200

if __name__ == '__main__':
    app.run(debug=True, port=5000)
