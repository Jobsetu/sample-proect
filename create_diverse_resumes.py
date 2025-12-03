#!/usr/bin/env python3
"""Create diverse test resumes in different formats to verify parsing robustness."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_traditional_resume():
    """Create a traditional chronological resume (TXT)."""
    content = """JANE SMITH
jane.smith@email.com | (555) 123-4567 | San Francisco, CA
LinkedIn: linkedin.com/in/janesmith | GitHub: github.com/janesmith

PROFESSIONAL SUMMARY
Senior Software Engineer with 5+ years of experience building scalable web applications and cloud infrastructure. Expert in Python, JavaScript, and AWS. Proven track record of leading teams and delivering high-impact projects.

TECHNICAL SKILLS
Languages: Python, JavaScript, TypeScript, Java, SQL
Frameworks: React, Node.js, Django, Flask, Express.js
Cloud & DevOps: AWS (EC2, S3, Lambda), Docker, Kubernetes, CI/CD, Terraform
Databases: PostgreSQL, MongoDB, Redis, MySQL
Tools: Git, JIRA, VS Code, Postman

PROFESSIONAL EXPERIENCE

Senior Software Engineer | TechStart Inc | San Francisco, CA | Jan 2021 - Present
- Led development of microservices architecture serving 2M+ daily active users
- Reduced API response time by 60% through optimization and caching strategies
- Mentored team of 5 junior engineers and conducted code reviews
- Implemented automated testing pipeline, increasing code coverage from 45% to 85%

Software Engineer | DataCorp | Remote | Jun 2019 - Dec 2020
- Developed RESTful APIs using Python/Flask for data analytics platform
- Built React-based dashboard for visualizing real-time metrics
- Collaborated with product team to define technical requirements
- Migrated legacy systems to cloud-based infrastructure on AWS

Junior Developer | StartupXYZ | Palo Alto, CA | Jul 2018 - May 2019
- Contributed to full-stack web development using React and Node.js
- Participated in daily standups and agile development practices
- Fixed bugs and implemented new features based on user feedback

EDUCATION

Bachelor of Science in Computer Science | Stanford University | 2018
GPA: 3.8/4.0 | Relevant Coursework: Algorithms, Data Structures, Machine Learning, Database Systems

CERTIFICATIONS
- AWS Certified Solutions Architect - Associate (2022)
- Certified Kubernetes Administrator - CKA (2021)
"""
    
    with open('test_resume_traditional.txt', 'w') as f:
        f.write(content)
    print("✓ Created test_resume_traditional.txt")

def create_functional_resume():
    """Create a skills-first functional resume (DOCX)."""
    doc = Document()
    
    # Header
    name = doc.add_paragraph('MICHAEL CHEN')
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name.runs[0].font.size = Pt(18)
    name.runs[0].font.bold = True
    
    contact = doc.add_paragraph('michael.chen@example.com | +1-650-555-0123 | Seattle, WA')
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.runs[0].font.size = Pt(10)
    
    # Core Competencies
    doc.add_heading('CORE COMPETENCIES', level=1)
    doc.add_paragraph('Full-Stack Development • Cloud Architecture • Agile/Scrum • Team Leadership • DevOps')
    
    # Technical Skills
    doc.add_heading('TECHNICAL EXPERTISE', level=1)
    skills_para = doc.add_paragraph()
    skills_para.add_run('Programming: ').bold = True
    skills_para.add_run('JavaScript, Python, Go, Ruby, C++\n')
    skills_para.add_run('Frontend: ').bold = True
    skills_para.add_run('React, Vue.js, Angular, HTML5, CSS3, Tailwind\n')
    skills_para.add_run('Backend: ').bold = True
    skills_para.add_run('Node.js, Express, Django, Rails, GraphQL\n')
    skills_para.add_run('Cloud: ').bold = True
    skills_para.add_run('AWS, Google Cloud, Azure, Serverless, Lambda\n')
    skills_para.add_run('DevOps: ').bold = True
    skills_para.add_run('Docker, Kubernetes, Jenkins, GitLab CI, Terraform')
    
    # Professional Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    
    exp1 = doc.add_paragraph()
    exp1.add_run('Lead Software Engineer | CloudTech Solutions | 2020 - Present\n').bold = True
    doc.add_paragraph('• Architected and deployed microservices handling 5M requests/day', style='List Bullet')
    doc.add_paragraph('• Led migration from monolith to microservices, reducing deployment time by 75%', style='List Bullet')
    doc.add_paragraph('• Managed team of 8 developers across 3 time zones', style='List Bullet')
    
    exp2 = doc.add_paragraph()
    exp2.add_run('Software Engineer | InnovateLabs | 2018 - 2020\n').bold = True
    doc.add_paragraph('• Built real-time data processing pipeline using Kafka and Spark', style='List Bullet')
    doc.add_paragraph('• Developed customer-facing dashboard using React and D3.js', style='List Bullet')
    doc.add_paragraph('• Reduced infrastructure costs by 40% through optimization', style='List Bullet')
    
    # Education
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('Master of Science in Computer Science | MIT | 2018')
    doc.add_paragraph('Bachelor of Science in Software Engineering | UC Berkeley | 2016')
    
    doc.save('test_resume_functional.docx')
    print("✓ Created test_resume_functional.docx")

def create_minimal_resume():
    """Create a minimal/sparse resume (TXT)."""
    content = """Bob Johnson
bob.j@mail.com
(415) 555-9876

Software Developer

SKILLS
Python, JavaScript, React, SQL, Git

EXPERIENCE

Software Developer - WebDev Co - 2022 to Present
Built web applications

Junior Developer - StartCo - 2020-2022
Worked on frontend

EDUCATION
BS Computer Science, State University, 2020
"""
    
    with open('test_resume_minimal.txt', 'w') as f:
        f.write(content)
    print("✓ Created test_resume_minimal.txt")

def create_dense_resume():
    """Create a dense/detailed resume (TXT formatted as PDF-like)."""
    content = """========================================
ALEXANDRA RODRIGUEZ
Senior Full-Stack Engineer & Tech Lead
========================================

Contact: alex.rodriguez@techmail.com | M: +1 (408) 555-TECH
Location: Austin, TX 78701 | LinkedIn: /in/alexrodriguez | GitHub: @alexrodriguez

────────────────────────────────────────
PROFESSIONAL SUMMARY
────────────────────────────────────────
Innovative and results-driven Senior Software Engineer with 8+ years of progressive experience in enterprise software development, cloud architecture, and technical leadership. Specializes in building highly scalable distributed systems, implementing DevOps best practices, and mentoring high-performing engineering teams. Passionate about leveraging cutting-edge technologies to solve complex business challenges and drive organizational growth.

────────────────────────────────────────
TECHNICAL PROFICIENCIES
────────────────────────────────────────
Languages & Frameworks:
  • Primary: JavaScript (ES6+), TypeScript, Python 3.x, Java, Go
  • Web: React.js, Next.js, Vue.js, Angular, Node.js, Express.js, FastAPI
  • Mobile: React Native, Flutter
  
Cloud & Infrastructure:
  • AWS: EC2, ECS, Lambda, S3, DynamoDB, RDS, CloudFormation, AWS CDK
  • GCP: Compute Engine, Cloud Functions, Cloud SQL, GKE
  • Azure: App Service, Functions, Cosmos DB
  
DevOps & Tools:
  • Container Orchestration: Docker, Kubernetes, Helm, Docker Compose
  • CI/CD: Jenkins, GitLab CI, GitHub Actions, CircleCI, ArgoCD
  • IaC: Terraform, Ansible, Pulumi
  • Monitoring: Datadog, Prometheus, Grafana, ELK Stack, New Relic
  
Data & Databases:
  • SQL: PostgreSQL, MySQL, Oracle, SQL Server
  • NoSQL: MongoDB, Cassandra, Redis, Elasticsearch
  • Stream Processing: Apache Kafka, RabbitMQ, AWS Kinesis

────────────────────────────────────────
PROFESSIONAL EXPERIENCE
────────────────────────────────────────

PRINCIPAL SOFTWARE ENGINEER & TECH LEAD
Enterprise Solutions Group | Austin, TX | March 2021 – Present

• Spearhead architecture and development of cloud-native SaaS platform serving Fortune 500 clients, processing 100M+ transactions monthly with 99.99% uptime SLA
• Led technical design and implementation of event-driven microservices architecture using Node.js, Python, and AWS, reducing system latency by 65% and improving throughput to 50K TPS
• Architected and deployed auto-scaling infrastructure on AWS ECS with Terraform, reducing operational costs by $2.3M annually while improving resource utilization from 42% to 87%
• Championed adoption of GitOps and CI/CD best practices, reducing deployment time from 4 hours to 12 minutes and increasing release frequency from monthly to multiple times daily
• Built and mentored cross-functional engineering team of 15 developers, establishing code quality standards, conducting architecture reviews, and fostering culture of continuous improvement
• Implemented comprehensive observability stack with Datadog and ELK, achieving 95% reduction in mean time to detection (MTTD) for production issues
• Drove migration from monolithic legacy system to microservices, completing 18-month project ahead of schedule and under budget with zero downtime

SENIOR SOFTWARE ENGINEER
TechInnovate Corp | Remote | June 2019 – February 2021

• Developed high-performance real-time analytics platform using React, Python, and PostgreSQL, enabling data-driven insights for 500K+ users
• Optimized database queries and implemented distributed caching with Redis, improving page load times by 78% and reducing database load by 60%
• Built RESTful and GraphQL APIs consumed by mobile and web applications, handling 2M+ API calls daily
• Implemented automated testing framework achieving 90%+ code coverage, reducing production bugs by 55%
• Collaborated with product and UX teams to deliver feature-rich customer dashboard with real-time data visualization using D3.js and WebSockets
• Led proof-of-concept for serverless architecture, resulting in 40% cost savings on compute resources

SOFTWARE ENGINEER II
DataStream Technologies | San Jose, CA | January 2017 – May 2019

• Engineered data ingestion pipelines processing 5TB+ daily using Apache Kafka, Spark, and Python
• Developed internal developer tools and CLI applications improving team productivity by 30%
• Contributed to open-source projects and internal libraries used across engineering organization
• Participated in on-call rotation, maintaining 99.9% service availability

────────────────────────────────────────
EDUCATION & CERTIFICATIONS
────────────────────────────────────────

Master of Science in Computer Science (AI Specialization)
Georgia Institute of Technology | GPA: 3.9/4.0 | 2020

Bachelor of Science in Software Engineering (Summa Cum Laude)
University of Texas at Austin | GPA: 3.85/4.0 | 2016

Certifications:
• AWS Certified Solutions Architect – Professional (2023)
• AWS Certified Developer – Associate (2022)
• Certified Kubernetes Administrator (CKA) - CNCF (2022)
• Google Cloud Professional Cloud Architect (2021)
• Certified Scrum Master (CSM) - Scrum Alliance (2020)

────────────────────────────────────────
PUBLICATIONS & SPEAKING
────────────────────────────────────────
• "Scaling Microservices: Lessons from Production" - AWS re:Invent 2022
• "Event-Driven Architecture Patterns" - Tech Blog (250K+ views)
• Conference Speaker at KubeCon, DevOps Days, React Conf
"""
    
    with open('test_resume_dense.txt', 'w') as f:
        f.write(content)
    print("✓ Created test_resume_dense.txt")

if __name__ == '__main__':
    print("Creating diverse test resumes...")
    print()
    
    create_traditional_resume()
    create_functional_resume()
    create_minimal_resume()
    create_dense_resume()
    
    print()
    print("✓ All test resumes created successfully!")
    print("\nFiles created:")
    print("  - test_resume_traditional.txt (traditional chronological format)")
    print("  - test_resume_functional.docx (skills-first functional format)")
    print("  - test_resume_minimal.txt (sparse/minimal information)")
    print("  - test_resume_dense.txt (detailed/comprehensive)")
